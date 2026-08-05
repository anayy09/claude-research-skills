#!/usr/bin/env python3
"""Inventory a publisher template so its structural requirements are explicit.

Accepts a zip, a directory, a single .tex/.cls file, or a .docx/.dotx template
and prints what the target document has to contain: document class and options,
front-matter macros, required environments and sections, bibliography style,
and the files that must ship with the submission.

Usage:
    python inspect_template.py sn-article-templates.zip
    python inspect_template.py ./elsarticle/ --json template.json
    python inspect_template.py journal-template.docx

Nothing is modified. Zip inputs are unpacked next to the archive unless
--extract-to is given.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path

FRONT_MACROS = [
    "title", "shorttitle", "titlerunning", "author", "authorrunning", "affil",
    "affiliation", "address", "email", "orcid", "date", "keywords", "keyword",
    "IEEEauthorblockN", "IEEEauthorblockA", "markboth", "IEEEtitleabstractindextext",
    "maketitle", "institute", "corres", "authormark", "presentaddress",
    "ead", "cortext", "fntext", "tnotetext", "abstracttext", "abstract",
    "bmhead", "backmatter", "frontmatter", "bibliography", "addbibresource",
    "Title", "Author", "AuthorNames", "AuthorCitation", "firstnote", "corres",
    "acmConference", "acmDOI", "acmISBN", "copyrightyear", "ccsdesc",
    "IEEEmembership", "IEEEpubid", "IEEEPARstart", "highlights", "shortauthors",
]

ENVIRONMENTS = [
    "abstract", "keywords", "IEEEkeywords", "frontmatter", "backmatter",
    "highlights", "graphicalabstract", "thebibliography", "acknowledgments",
    "acknowledgements", "table", "figure", "algorithm", "theorem", "proof",
    "appendices", "appendix", "sidewaystable", "landscape",
]

BIB_HINTS = [
    (r"\\bibliographystyle\{([^}]*)\}", "bibliographystyle"),
    (r"\\bibliography\{([^}]*)\}", "bibliography"),
    (r"\\usepackage(?:\[[^\]]*\])?\{(natbib|biblatex|cite|apacite)\}", "bib package"),
    (r"\\addbibresource\{([^}]*)\}", "biblatex resource"),
]

TEXT_EXT = {".tex", ".cls", ".sty", ".bst", ".bib", ".txt", ".md", ".cfg", ".ins", ".dtx"}


def unpack(path: Path, extract_to: Path | None) -> Path:
    dest = extract_to or path.with_suffix("")
    dest.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        for member in zf.namelist():
            target = (dest / member).resolve()
            if not str(target).startswith(str(dest.resolve())):
                raise SystemExit("refusing to extract outside destination: %s" % member)
        zf.extractall(dest)
    return dest


def find_main_tex(files: list) -> Path | None:
    candidates = [f for f in files if f.suffix.lower() == ".tex"]
    scored = []
    for f in candidates:
        try:
            text = f.read_text(encoding="utf-8", errors="replace")
        except Exception:
            continue
        score = 0
        if "\\documentclass" in text:
            score += 10
        if "\\begin{document}" in text:
            score += 10
        score += min(len(text) // 2000, 10)
        if re.search(r"(sample|template|main|manuscript|article)", f.name, re.I):
            score += 5
        scored.append((score, len(text), f))
    if not scored:
        return None
    scored.sort(key=lambda t: (t[0], t[1]), reverse=True)
    return scored[0][2]


def analyze_tex(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="replace")
    body = text.split("\\begin{document}", 1)
    preamble = body[0]
    doc = body[1] if len(body) > 1 else ""

    m = re.search(r"\\documentclass\s*(\[[^\]]*\])?\s*\{([^}]*)\}", text)
    info = {
        "file": str(path),
        "documentclass": m.group(2) if m else None,
        "class_options": [o.strip() for o in (m.group(1) or "[]")[1:-1].split(",") if o.strip()],
        "packages": sorted(set(re.findall(r"\\usepackage(?:\[[^\]]*\])?\{([^}]*)\}", preamble))),
        "front_matter_macros": [],
        "environments": [],
        "sections": re.findall(r"\\(?:sub)*section\*?\{([^}]*)\}", doc),
        "bibliography": {},
        "labels": re.findall(r"\\label\{([^}]*)\}", doc),
        "custom_commands": sorted(set(re.findall(r"\\(?:new|renew|provide)command\*?\{?\\(\w+)", preamble))),
    }
    for macro in FRONT_MACROS:
        if re.search(r"\\%s\b" % re.escape(macro), text):
            info["front_matter_macros"].append(macro)
    for env in ENVIRONMENTS:
        if re.search(r"\\begin\{%s\*?\}" % re.escape(env), text):
            info["environments"].append(env)
    for pat, label in BIB_HINTS:
        found = re.findall(pat, text)
        if found:
            info["bibliography"][label] = sorted(set(found))
    return info


def analyze_docx(path: Path) -> dict:
    info = {"file": str(path), "styles": [], "headings": [], "page_setup": {}}
    try:
        import docx  # type: ignore
    except ImportError:
        info["error"] = "python-docx not installed; cannot inspect .docx styles"
        return info
    d = docx.Document(str(path))
    info["styles"] = sorted({s.name for s in d.styles})
    for p in d.paragraphs:
        name = (p.style.name or "").lower()
        if name.startswith("heading") or name in ("title", "subtitle"):
            if p.text.strip():
                info["headings"].append({"style": p.style.name, "text": p.text.strip()})
    if d.sections:
        s = d.sections[0]
        def pts(v):
            return round(v.pt, 1) if v is not None else None
        info["page_setup"] = {
            "page_width_pt": pts(s.page_width), "page_height_pt": pts(s.page_height),
            "left_margin_pt": pts(s.left_margin), "right_margin_pt": pts(s.right_margin),
            "top_margin_pt": pts(s.top_margin), "bottom_margin_pt": pts(s.bottom_margin),
        }
    info["tables"] = len(d.tables)
    info["paragraphs"] = len(d.paragraphs)
    return info


def classify(files: list) -> dict:
    groups = {"class": [], "style": [], "bibstyle": [], "bibdata": [], "tex": [],
              "docx": [], "figures": [], "docs": [], "other": []}
    for f in files:
        ext = f.suffix.lower()
        if ext == ".cls":
            groups["class"].append(f)
        elif ext == ".sty":
            groups["style"].append(f)
        elif ext == ".bst":
            groups["bibstyle"].append(f)
        elif ext == ".bib":
            groups["bibdata"].append(f)
        elif ext == ".tex":
            groups["tex"].append(f)
        elif ext in (".docx", ".dotx", ".doc"):
            groups["docx"].append(f)
        elif ext in (".pdf", ".png", ".jpg", ".jpeg", ".eps", ".tif", ".tiff"):
            groups["figures"].append(f)
        elif ext in (".txt", ".md", ".rst") or f.name.lower().startswith("readme"):
            groups["docs"].append(f)
        else:
            groups["other"].append(f)
    return groups


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("path")
    ap.add_argument("--extract-to")
    ap.add_argument("--json", dest="json_out")
    args = ap.parse_args()

    path = Path(args.path)
    if not path.exists():
        print("not found: %s" % path, file=sys.stderr)
        return 2

    root = path
    if path.is_file() and path.suffix.lower() == ".zip":
        root = unpack(path, Path(args.extract_to) if args.extract_to else None)
        print("unpacked to %s" % root)

    if root.is_file():
        files = [root]
    else:
        files = [f for f in sorted(root.rglob("*"))
                 if f.is_file() and "__MACOSX" not in str(f)]

    groups = classify(files)
    result = {"root": str(root), "files": {k: [str(p) for p in v] for k, v in groups.items() if v}}

    main_tex = find_main_tex(files)
    if main_tex:
        result["latex"] = analyze_tex(main_tex)
    docx_files = groups["docx"]
    if docx_files and not main_tex:
        result["docx"] = analyze_docx(docx_files[0])
    elif docx_files:
        result["docx"] = {"file": str(docx_files[0]), "note": "LaTeX template also present"}

    readmes = groups["docs"]
    if readmes:
        result["instructions_files"] = [str(p) for p in readmes]

    # ---- human-readable output ----
    print("\n=== template inventory: %s ===" % root)
    for key in ("class", "style", "bibstyle", "bibdata", "tex", "docx", "docs"):
        if groups[key]:
            print("%-9s %s" % (key + ":", ", ".join(p.name for p in groups[key][:12])))
    if groups["figures"]:
        print("%-9s %d files" % ("figures:", len(groups["figures"])))

    tex = result.get("latex")
    if tex:
        print("\nmain .tex        %s" % Path(tex["file"]).name)
        print("documentclass    %s %s" % (tex["documentclass"],
                                          "[" + ",".join(tex["class_options"]) + "]"
                                          if tex["class_options"] else ""))
        cls_present = any(p.stem == tex["documentclass"] for p in groups["class"])
        print("class file       %s" % ("bundled in template" if cls_present
                                       else "NOT bundled: must be installed or fetched"))
        print("front matter     %s" % ", ".join(tex["front_matter_macros"]) or "none detected")
        print("environments     %s" % ", ".join(tex["environments"]) or "none detected")
        if tex["bibliography"]:
            for k, v in tex["bibliography"].items():
                print("%-16s %s" % (k, ", ".join(v)))
        else:
            print("bibliography     none declared in the sample")
        if tex["sections"]:
            print("sample sections  %s" % " | ".join(tex["sections"][:15]))
        if tex["custom_commands"]:
            print("custom macros    %s" % ", ".join(tex["custom_commands"][:15]))

    dx = result.get("docx")
    if dx and "styles" in dx:
        print("\ndocx styles      %s" % ", ".join(dx["styles"][:20]))
        if dx.get("headings"):
            print("docx headings    %s" % " | ".join(h["text"] for h in dx["headings"][:15]))
        if dx.get("page_setup"):
            print("page setup       %s" % dx["page_setup"])

    if readmes:
        print("\nread these for rules the template does not encode:")
        for p in readmes[:6]:
            print("  %s" % p)

    print("\nnext: record class, options, section order, citation style, figure "
          "format, and length limits in the requirements sheet, then confirm them "
          "against the journal's author instructions page.")

    if args.json_out:
        Path(args.json_out).write_text(json.dumps(result, indent=2, ensure_ascii=False),
                                       encoding="utf-8")
        print("wrote %s" % args.json_out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
